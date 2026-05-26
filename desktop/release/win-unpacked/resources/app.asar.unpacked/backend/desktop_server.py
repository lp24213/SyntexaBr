#!/usr/bin/env python3
"""
SYNTEXA DESKTOP BACKEND SERVER V45
====================================
Servidor HTTP local para o app desktop Electron.
Expõe Foundation Model, multimodal, exports, e inferência 70B via REST.
Boot Validation Obrigatória — PROIBIDO subir UI com modelo quebrado.

Uso:
    python desktop_server.py --port 34560
"""
from __future__ import annotations

import argparse
import base64
import hashlib
import json
import logging
import os
import sys
import tempfile
import time
import traceback
import unicodedata
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

# Paths relativos ao desktop/backend/
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ── LOGS ENTERPRISE ───────────────────────────────────────
LOGS_DIR = Path(os.environ.get("SYNTEXA_LOGS_DIR", str(ROOT / "logs")))
LOGS_DIR.mkdir(parents=True, exist_ok=True)

# Handlers separados por domínio
_log_format = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s — %(message)s")

for _log_name, _fname in [
    ("inference", "inference.log"),
    ("tokenizer", "tokenizer.log"),
    ("runtime", "runtime.log"),
    ("cuda", "cuda.log"),
    ("streaming", "streaming.log"),
    ("multimodal", "multimodal.log"),
    ("quantum", "quantum.log"),
]:
    _lh = logging.FileHandler(LOGS_DIR / _fname, encoding="utf-8")
    _lh.setFormatter(_log_format)
    logging.getLogger(_log_name).addHandler(_lh)
    logging.getLogger(_log_name).setLevel(logging.INFO)

def _enterprise_log(domain: str, level: str, msg: str, **extra: Any) -> None:
    entry = {
        "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S") + f".{int(time.time()*1000)%1000:03d}",
        "domain": domain,
        "level": level,
        "message": msg,
        **extra,
    }
    path = LOGS_DIR / f"{domain}.log"
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    getattr(logging.getLogger(domain), level.lower(), log.info)("[%s] %s", domain.upper(), msg)

# ── DEPENDÊNCIAS OPCIONAIS ────────────────────────────────
try:
    import uvicorn
    from fastapi import FastAPI, HTTPException, Request
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import StreamingResponse, JSONResponse, PlainTextResponse
    from pydantic import BaseModel, Field
    HAS_FASTAPI = True
except ImportError:
    HAS_FASTAPI = False
    log.error("fastapi/uvicorn não instalados. Instale: pip install fastapi uvicorn")

# ── MODELOS PYDANTIC ──────────────────────────────────────

class ChatRequest(BaseModel):
    messages: List[Dict[str, str]]
    max_new_tokens: int = Field(default=512, ge=1, le=32768)
    temperature: float = Field(default=0.7, ge=0.0, le=2.0)
    top_p: float = Field(default=0.9, ge=0.0, le=1.0)
    stream: bool = True
    tensor_parallel_size: Optional[int] = Field(default=None, ge=1)
    quantization: Optional[str] = Field(default=None, pattern="^(int4|int8|fp16|bf16)$")
    use_flash_attention: bool = True
    use_paged_attention: bool = True


class MultimodalRequest(BaseModel):
    base64_data: str
    mime_type: str
    kind: str = "image"  # image, audio, pdf, docx


class ExportRequest(BaseModel):
    messages: List[Dict[str, Any]]
    format: str = "md"  # md, html, csv, json, txt, pdf, docx


class RuntimeStatus(BaseModel):
    runtime_ready: bool
    model_loaded: bool
    model_size_b: Optional[int] = None
    device: str = "cpu"
    cuda_available: bool = False
    tensor_parallel_size: int = 1
    quantization: Optional[str] = None
    version: str = "V45"


# ── APP FASTAPI ───────────────────────────────────────────
app = FastAPI(title="Syntexa Desktop Runtime", version="V45")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Estado do runtime
_runtime_ready = False
_engine = None
_runtime_status = RuntimeStatus(runtime_ready=False, model_loaded=False)
_boot_validator = None  # type: ignore
_boot_report = {}  # type: ignore

# ── HELPERS ───────────────────────────────────────────────

def _sanitize(text: str) -> str:
    """Remove caracteres de controle exceto newline e tab."""
    return "".join(ch for ch in text if ch == "\n" or ch == "\t" or unicodedata.category(ch)[0] != "C")


def _fail_fast(response_text: str, context: str = "") -> str:
    """Fail fast: se resposta vazia ou só whitespace, levanta erro real."""
    if not response_text or not str(response_text).strip():
        raise RuntimeError(
            f"[Syntexa V43] Inferência retornou resposta vazia. "
            f"Contexto: {context or 'chat completion'}. "
            f"Verifique disponibilidade do runtime LLM local e logs do backend."
        )
    return str(response_text).strip()


def _get_cuda_info() -> dict:
    """Detecta CUDA e GPUs disponíveis."""
    try:
        import torch
        return {
            "available": torch.cuda.is_available(),
            "device_count": torch.cuda.device_count(),
            "device_name": torch.cuda.get_device_name(0) if torch.cuda.is_available() else None,
        }
    except Exception:
        return {"available": False, "device_count": 0, "device_name": None}


def _get_boot_failure_detail() -> str:
    """Retorna diagnóstico técnico real quando o runtime não está pronto."""
    if _boot_report and _boot_report.get("failures"):
        failures = _boot_report["failures"]
        msgs = [f"{f['component']}: {f['error']}" for f in failures[:3]]
        return f"[Syntexa V45] BOOT BLOQUEADO — {len(failures)} falha(s). " + " | ".join(msgs)
    return "[Syntexa V45] Runtime não pronto. Nenhum modelo carregado ou boot validation falhou."


# ── STARTUP / MODEL LOADING ───────────────────────────────

@app.on_event("startup")
async def startup():
    global _runtime_ready, _engine, _runtime_status, _boot_validator, _boot_report
    log.info("[DESKTOP SERVER V45] Iniciando runtime soberano...")
    _enterprise_log("runtime", "info", "Startup iniciado", version="V45")

    cuda = _get_cuda_info()
    _runtime_status.cuda_available = cuda["available"]
    _runtime_status.device = "cuda" if cuda["available"] else "cpu"
    _enterprise_log("cuda", "info", "CUDA info", **cuda)

    # ── CARREGAMENTO DO ENGINE ──────────────────────────────
    engine_loaded = False
    load_errors = []

    # 1) Tenta carregar Foundation Model própria
    try:
        from vereda_ai.syntexa_core.foundation_runtime import get_foundation_runtime
        rt = get_foundation_runtime(checkpoint_dir=str(ROOT / "checkpoints" / "foundation"))
        loaded = rt.load()
        if loaded:
            _engine = rt.engine
            engine_loaded = True
            _enterprise_log("runtime", "info", "Foundation Model carregada")
        else:
            load_errors.append("Foundation Model: checkpoint inexistente ou não treinado.")
            _enterprise_log("runtime", "warning", "Foundation Model não carregada")
    except Exception as e:
        load_errors.append(f"Foundation Model: {type(e).__name__}: {e}")
        _enterprise_log("runtime", "error", f"Foundation runtime falhou: {e}", traceback=traceback.format_exc())

    # 2) Fallback: llama-cpp-python para GGUF local
    if not engine_loaded:
        gguf_path = os.environ.get("SYNTEXA_GGUF_PATH", "")
        if gguf_path and Path(gguf_path).is_file():
            try:
                from llama_cpp import Llama
                n_gpu_layers = -1 if cuda["available"] else 0
                _engine = Llama(
                    model_path=gguf_path,
                    n_ctx=8192,
                    n_gpu_layers=n_gpu_layers,
                    verbose=False,
                )
                engine_loaded = True
                _enterprise_log("runtime", "info", f"llama.cpp GGUF carregado: {gguf_path}")
            except Exception as e:
                load_errors.append(f"llama.cpp: {type(e).__name__}: {e}")
                _enterprise_log("runtime", "error", f"llama.cpp falhou: {e}")

    # 3) Fallback: transformers AutoModel local
    if not engine_loaded:
        local_model = os.environ.get("SYNTEXA_LOCAL_MODEL", "")
        if local_model:
            try:
                from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
                tok = AutoTokenizer.from_pretrained(local_model, local_files_only=True)
                model = AutoModelForCausalLM.from_pretrained(
                    local_model,
                    local_files_only=True,
                    torch_dtype="auto",
                    device_map="auto" if cuda["available"] else None,
                )
                _engine = pipeline("text-generation", model=model, tokenizer=tok, device=0 if cuda["available"] else -1)
                engine_loaded = True
                _enterprise_log("runtime", "info", f"transformers local carregado: {local_model}")
            except Exception as e:
                load_errors.append(f"transformers: {type(e).__name__}: {e}")
                _enterprise_log("runtime", "error", f"transformers local falhou: {e}")

    # ── BOOT VALIDATION OBRIGATÓRIA ─────────────────────────
    if engine_loaded:
        try:
            from desktop.backend.boot_validator import run_boot_validation
            _boot_validator = run_boot_validation(_engine, checkpoint_dir=str(ROOT / "checkpoints" / "foundation"))
            _boot_report = _boot_validator.get_diagnostic_report()
            if _boot_validator.is_bootable():
                _runtime_ready = True
                _runtime_status.model_loaded = True
                _runtime_status.runtime_ready = True
                _enterprise_log("runtime", "info", "BOOT VALIDATION APROVADO — RUNTIME READY")
                log.info("[DESKTOP SERVER] RUNTIME READY (Foundation Model V45)")
            else:
                _runtime_ready = False
                _runtime_status.runtime_ready = False
                failed_names = [f["component"] for f in _boot_report.get("failures", [])]
                log.error("[DESKTOP SERVER] BOOT BLOQUEADO: %s", ", ".join(failed_names))
                _enterprise_log("runtime", "error", "BOOT BLOQUEADO", failures=failed_names)
        except Exception as e:
            log.error("[DESKTOP SERVER] Boot validator falhou: %s", e)
            _enterprise_log("runtime", "error", f"Boot validator exceção: {e}", traceback=traceback.format_exc())
            _runtime_ready = False
            _runtime_status.runtime_ready = False
    else:
        _runtime_ready = False
        _runtime_status.runtime_ready = False
        log.error("[DESKTOP SERVER] NENHUM ENGINE CARREGADO.")
        _enterprise_log("runtime", "error", "Nenhum engine carregado", errors=load_errors)


# ── HEALTH ────────────────────────────────────────────────

@app.get("/health")
def health():
    return {
        "status": "ok" if _runtime_ready else "degraded",
        "runtime_ready": _runtime_ready,
        "model_loaded": _runtime_status.model_loaded,
        "device": _runtime_status.device,
        "cuda_available": _runtime_status.cuda_available,
        "version": "V45",
        "mode": "soberano-offline",
        "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "boot_report": _boot_report if _boot_report else None,
    }


@app.get("/boot/diagnostic")
def boot_diagnostic():
    """
    Diagnóstico técnico real.
    NUNCA retorna 'tente novamente'. Retorna dados reais do boot.
    """
    if not _boot_report:
        return JSONResponse(
            status_code=503,
            content={
                "boot_passed": False,
                "version": "V45",
                "error": "Boot validator não executado. Engine não carregado.",
                "components": {"engine": "none", "checkpoint": "not_found"},
            },
        )
    return _boot_report


@app.get("/runtime/status")
def runtime_status():
    return _runtime_status.model_dump()


# ── CHAT ──────────────────────────────────────────────────

@app.post("/chat")
def chat(req: ChatRequest):
    if not _runtime_ready:
        _enterprise_log("inference", "error", "Chat bloqueado: runtime não pronto")
        detail = _get_boot_failure_detail()
        raise HTTPException(status_code=503, detail=detail)
    start = time.time()
    try:
        response = _compute_chat(req)
        tokens = _approx_tokens(response)
        latency = time.time() - start
        _enterprise_log("inference", "info", "Chat completion OK", tokens=tokens, latency_ms=round(latency*1000, 1))
        log.info("[CHAT] %.2fs | %d tokens | temp=%.2f", latency, tokens, req.temperature)
        return {"response": response, "tokens": tokens, "time_sec": round(latency, 3)}
    except Exception as e:
        _enterprise_log("inference", "error", "Chat erro", error=str(e), traceback=traceback.format_exc())
        log.error("[CHAT] Erro: %s", traceback.format_exc())
        # FAIL FAST ABSOLUTO: retorna erro técnico real, nunca genérico
        raise HTTPException(status_code=500, detail=f"[Syntexa V45] Inferência falhou: {type(e).__name__}: {e}")


@app.post("/chat/stream")
def chat_stream(req: ChatRequest):
    if not _runtime_ready:
        detail = _get_boot_failure_detail()
        _enterprise_log("streaming", "error", "Stream bloqueado: runtime não pronto")
        raise HTTPException(status_code=503, detail=detail)

    def generate():
        t0 = time.time()
        chunk_count = 0
        try:
            for chunk in _compute_chat_stream(req):
                if chunk:
                    chunk_count += 1
                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
            latency = time.time() - t0
            _enterprise_log("streaming", "info", "Stream concluído", chunks=chunk_count, latency_ms=round(latency*1000, 1))
        except Exception as e:
            _enterprise_log("streaming", "error", "Stream erro", error=str(e), traceback=traceback.format_exc())
            log.error("[STREAM] Erro: %s", e)
            # FAIL FAST ABSOLUTO
            yield f"data: {json.dumps({'error': f'[Syntexa V45] Stream falhou: {type(e).__name__}: {e}'})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── CORE GENERATION (fail-fast) ────────────────────────────

def _compute_chat(req: ChatRequest) -> str:
    """Geração não-streaming com fail-fast real."""
    messages = req.messages
    if not messages:
        raise RuntimeError("Nenhuma mensagem fornecida.")

    engine_type = _detect_engine_type()

    if engine_type == "foundation":
        out = _engine.chat(
            messages=messages,
            max_new_tokens=req.max_new_tokens,
            temperature=req.temperature,
            top_p=req.top_p,
        )
    elif engine_type == "llama_cpp":
        prompt = _messages_to_prompt(messages)
        result = _engine(prompt, max_tokens=req.max_new_tokens, temperature=req.temperature, top_p=req.top_p)
        out = result["choices"][0]["text"] if result and "choices" in result else ""
    elif engine_type == "transformers":
        prompt = _messages_to_prompt(messages)
        result = _engine(prompt, max_new_tokens=req.max_new_tokens, temperature=req.temperature, do_sample=True)
        out = result[0]["generated_text"] if result else ""
        if out.startswith(prompt):
            out = out[len(prompt):]
    else:
        raise RuntimeError("Engine desconhecido ou não carregado.")

    return _fail_fast(_sanitize(out), "chat completion")


def _compute_chat_stream(req: ChatRequest) -> Iterator[str]:
    """Streaming token por token com fail-fast."""
    messages = req.messages
    if not messages:
        raise RuntimeError("Nenhuma mensagem fornecida.")

    engine_type = _detect_engine_type()

    if engine_type == "foundation":
        for chunk in _engine.chat_stream(
            messages=messages,
            max_new_tokens=req.max_new_tokens,
            temperature=req.temperature,
            top_p=req.top_p,
        ):
            if chunk:
                yield chunk
    elif engine_type == "llama_cpp":
        prompt = _messages_to_prompt(messages)
        stream = _engine(prompt, max_tokens=req.max_new_tokens, temperature=req.temperature, stream=True)
        for item in stream:
            chunk = item["choices"][0]["text"] if item and "choices" in item else ""
            if chunk:
                yield chunk
    elif engine_type == "transformers":
        out = _compute_chat(req)
        words = out.split(" ")
        for w in words:
            yield w + " "
    else:
        raise RuntimeError("Engine desconhecido ou não carregado.")


def _detect_engine_type() -> str:
    if _engine is None:
        return "none"
    mod = type(_engine).__module__
    name = type(_engine).__name__
    if "foundation" in mod or "SyntexaInferenceEngine" in name:
        return "foundation"
    if "llama_cpp" in mod or "Llama" in name:
        return "llama_cpp"
    if "transformers" in mod or "Pipeline" in name:
        return "transformers"
    return "unknown"


def _messages_to_prompt(messages: list) -> str:
    """Converte mensagens ChatML para prompt texto simples."""
    parts = []
    for m in messages:
        role = (m.get("role") or "user").lower()
        content = (m.get("content") or "").strip()
        if role == "system":
            parts.append(f"<|system|>\n{content}")
        elif role == "user":
            parts.append(f"<|user|>\n{content}")
        elif role == "assistant":
            parts.append(f"<|assistant|>\n{content}")
    parts.append("<|assistant|>\n")
    return "\n".join(parts)


def _approx_tokens(text: str) -> int:
    """Estima tokens aproximados."""
    return max(1, int(len(text) * 0.75))


# ── MULTIMODAL ────────────────────────────────────────────

@app.post("/multimodal/process")
def multimodal_process(req: MultimodalRequest):
    """Processa upload multimodal (OCR, STT, vision, PDF, DOCX)."""
    try:
        data = base64.b64decode(req.base64_data)
    except Exception as e:
        _enterprise_log("multimodal", "error", "Base64 inválido", error=str(e))
        raise HTTPException(status_code=400, detail=f"Base64 inválido: {e}")

    t0 = time.time()
    if req.kind == "image":
        result = _process_image(data, req.mime_type)
    elif req.kind == "pdf":
        result = _process_pdf(data)
    elif req.kind == "audio":
        result = _process_audio(data, req.mime_type)
    elif req.kind == "docx":
        result = _process_docx(data)
    else:
        result = {"type": req.kind, "status": "unknown", "error": "Tipo não suportado"}

    latency = time.time() - t0
    _enterprise_log("multimodal", "info" if result.get("status") != "error" else "error",
                    f"Multimodal {req.kind}", status=result.get("status"), latency_ms=round(latency*1000, 1))
    return result


def _process_image(data: bytes, mime: str) -> dict:
    try:
        from vereda_ai.syntexa_core.multimodal.vision import SyntexaVisionEncoder
        enc = SyntexaVisionEncoder()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(data)
            tmp = f.name
        embedding = enc.encode_image(tmp)
        Path(tmp).unlink(missing_ok=True)
        return {"type": "image", "embedding_dim": len(embedding), "status": "encoded"}
    except Exception as e:
        log.warning("[VISION] Falha: %s", e)
        return {"type": "image", "error": str(e), "status": "error"}


def _process_pdf(data: bytes) -> dict:
    try:
        from vereda_ai.syntexa_core.multimodal.ocr import SyntexaOCR
        ocr = SyntexaOCR()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(data)
            tmp = f.name
        text = ocr.extract_text_from_pdf(tmp)
        Path(tmp).unlink(missing_ok=True)
        return {"type": "pdf", "text": text[:8000], "status": "extracted", "length": len(text)}
    except Exception as e:
        log.warning("[OCR] Falha: %s", e)
        try:
            import fitz
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(data)
                tmp = f.name
            doc = fitz.open(tmp)
            texts = [page.get_text() for page in doc]
            doc.close()
            Path(tmp).unlink(missing_ok=True)
            full = "\n".join(texts)
            return {"type": "pdf", "text": full[:8000], "status": "extracted_fallback", "length": len(full)}
        except Exception as e2:
            return {"type": "pdf", "error": f"OCR: {e}; PyMuPDF: {e2}", "status": "error"}


def _process_audio(data: bytes, mime: str) -> dict:
    try:
        from vereda_ai.syntexa_core.multimodal.stt import SyntexaSTT
        stt = SyntexaSTT()
        ext = ".wav" if "wav" in mime else ".mp3" if "mp3" in mime else ".ogg"
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
            f.write(data)
            tmp = f.name
        text = stt.transcribe(tmp)
        Path(tmp).unlink(missing_ok=True)
        return {"type": "audio", "transcript": text, "status": "transcribed"}
    except Exception as e:
        log.warning("[STT] Falha: %s", e)
        return {"type": "audio", "error": str(e), "status": "error"}


def _process_docx(data: bytes) -> dict:
    try:
        import docx
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(data)
            tmp = f.name
        document = docx.Document(tmp)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        Path(tmp).unlink(missing_ok=True)
        text = "\n".join(paragraphs)
        return {"type": "docx", "text": text[:8000], "status": "extracted", "length": len(text)}
    except Exception as e:
        return {"type": "docx", "error": str(e), "status": "error"}


# ── TTS ───────────────────────────────────────────────────

@app.post("/tts/synthesize")
def tts_synthesize(req: Dict[str, Any]):
    text = req.get("text", "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Texto não fornecido.")
    try:
        from vereda_ai.syntexa_core.multimodal.tts import SyntexaTTS
        tts = SyntexaTTS()
        import tempfile as _tf
        with _tf.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            out = tts.synthesize(text, output_path=f.name)
        with open(out, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        Path(out).unlink(missing_ok=True)
        _enterprise_log("multimodal", "info", "TTS OK", chars=len(text))
        return {"audio_base64": b64, "mime": "audio/wav", "status": "synthesized", "length_chars": len(text)}
    except Exception as e:
        _enterprise_log("multimodal", "error", "TTS falhou", error=str(e))
        log.error("[TTS] Falha: %s", e)
        raise HTTPException(status_code=500, detail=f"TTS falhou: {type(e).__name__}: {e}")


# ── EXPORT ────────────────────────────────────────────────

@app.post("/export")
def export_conversation(req: ExportRequest):
    _enterprise_log("runtime", "info", "Export iniciado", format=req.format, messages=len(req.messages))
    """Exporta conversa para formato solicitado com timestamps e metadata."""
    visible = [m for m in req.messages if m.get("role") != "system"]
    fmt = req.format.lower()

    if fmt == "txt":
        lines = ["Conversa Syntexa — Exportado em " + time.strftime("%Y-%m-%d %H:%M:%S")]
        lines.append("=" * 50)
        for m in visible:
            role = "Você" if m.get("role") == "user" else "Syntexa"
            ts = m.get("timestamp", "")
            lines.append(f"[{ts}] {role}")
            lines.append(m.get("content", ""))
            lines.append("")
        content = "\n".join(lines)
        mime = "text/plain"

    elif fmt == "md":
        lines = ["# Conversa Syntexa", ""]
        lines.append(f"_Exportado em {time.strftime('%Y-%m-%d %H:%M:%S')}_")
        lines.append("")
        for m in visible:
            role = "**Você**" if m.get("role") == "user" else "**Syntexa**"
            ts = m.get("timestamp", "")
            lines.append(f"## {role} — {ts}")
            lines.append("")
            lines.append(m.get("content", ""))
            lines.append("")
            lines.append("---")
            lines.append("")
        content = "\n".join(lines)
        mime = "text/markdown"

    elif fmt == "html":
        parts = [
            "<!DOCTYPE html><html><head>",
            "<meta charset='utf-8'><title>Syntexa Export</title>",
            "<style>body{font-family:system-ui,sans-serif;max-width:720px;margin:40px auto;line-height:1.6;color:#1a1a1a}"
            ".msg{margin:16px 0;padding:16px;border-radius:8px}"
            ".user{background:#f3f4f6}"
            ".assistant{background:#eef2ff}"
            ".role{font-weight:700;font-size:12px;text-transform:uppercase;letter-spacing:0.05em;margin-bottom:8px;color:#666}"
            "</style></head><body>",
            f"<h1>Conversa Syntexa</h1><p><i>Exportado em {time.strftime('%Y-%m-%d %H:%M:%S')}</i></p>",
        ]
        for m in visible:
            role = m.get("role", "user")
            cls = "user" if role == "user" else "assistant"
            label = "Você" if role == "user" else "Syntexa"
            ts = m.get("timestamp", "")
            parts.append(f'<div class="msg {cls}"><div class="role">{label} — {ts}</div><div>{_html_escape(m.get("content",""))}</div></div>')
        parts.append("</body></html>")
        content = "\n".join(parts)
        mime = "text/html"

    elif fmt == "csv":
        lines = ["timestamp,role,content,message_id"]
        for idx, m in enumerate(visible):
            ts = m.get("timestamp", "")
            role = m.get("role", "")
            content_escaped = (m.get("content", "") or "").replace('"', '""').replace("\n", " ")
            lines.append(f'{ts},{role},"{content_escaped}",{idx}')
        content = "\n".join(lines)
        mime = "text/csv"

    elif fmt == "json":
        payload = {
            "export_version": "V43",
            "exported_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
            "message_count": len(visible),
            "messages": visible,
        }
        content = json.dumps(payload, ensure_ascii=False, indent=2)
        mime = "application/json"

    elif fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas as pdfcanvas
            import io
            buf = io.BytesIO()
            c = pdfcanvas.Canvas(buf, pagesize=A4)
            w, h = A4
            y = h - 60
            c.setFont("Helvetica-Bold", 16)
            c.drawString(40, y, "Conversa Syntexa")
            y -= 24
            c.setFont("Helvetica", 10)
            c.drawString(40, y, f"Exportado em {time.strftime('%Y-%m-%d %H:%M:%S')}")
            y -= 30
            c.setFont("Helvetica", 11)
            for m in visible:
                role = "Você" if m.get("role") == "user" else "Syntexa"
                ts = m.get("timestamp", "")
                text = f"[{ts}] {role}: {m.get('content', '')}"
                words = text.split(" ")
                line = ""
                for word in words:
                    if c.stringWidth(line + " " + word, "Helvetica", 11) > w - 80:
                        c.drawString(40, y, line)
                        y -= 16
                        line = word
                        if y < 60:
                            c.showPage()
                            y = h - 60
                    else:
                        line = line + " " + word if line else word
                if line:
                    c.drawString(40, y, line)
                    y -= 16
                y -= 8
                if y < 60:
                    c.showPage()
                    y = h - 60
            c.save()
            buf.seek(0)
            b64 = base64.b64encode(buf.read()).decode()
            return {"base64": b64, "mime": "application/pdf", "status": "generated"}
        except Exception as e:
            log.warning("[EXPORT PDF] ReportLab falhou: %s. Retornando HTML fallback.", e)
            return export_conversation(ExportRequest(messages=req.messages, format="html"))

    elif fmt == "docx":
        try:
            import docx
            document = docx.Document()
            document.add_heading("Conversa Syntexa", level=0)
            document.add_paragraph(f"Exportado em {time.strftime('%Y-%m-%d %H:%M:%S')}")
            for m in visible:
                role = "Você" if m.get("role") == "user" else "Syntexa"
                ts = m.get("timestamp", "")
                p = document.add_paragraph()
                p.add_run(f"[{ts}] {role}").bold = True
                document.add_paragraph(m.get("content", ""))
            buf = io.BytesIO()
            document.save(buf)
            buf.seek(0)
            b64 = base64.b64encode(buf.read()).decode()
            return {"base64": b64, "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "status": "generated"}
        except Exception as e:
            log.warning("[EXPORT DOCX] python-docx falhou: %s", e)
            return export_conversation(ExportRequest(messages=req.messages, format="md"))

    else:
        raise HTTPException(status_code=400, detail=f"Formato não suportado: {fmt}")

    return {"content": content, "mime": mime, "status": "generated", "format": fmt}


def _html_escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ── MAIN ──────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--port", type=int, default=34560)
    ap.add_argument("--host", default="127.0.0.1")
    ap.add_argument("--gguf", default=os.environ.get("SYNTEXA_GGUF_PATH", ""), help="Caminho para modelo GGUF (llama.cpp)")
    ap.add_argument("--model", default=os.environ.get("SYNTEXA_LOCAL_MODEL", ""), help="Caminho/nome modelo transformers")
    args = ap.parse_args()

    if args.gguf:
        os.environ["SYNTEXA_GGUF_PATH"] = args.gguf
    if args.model:
        os.environ["SYNTEXA_LOCAL_MODEL"] = args.model

    if not HAS_FASTAPI:
        print("ERROR: fastapi/uvicorn não instalados")
        sys.exit(1)

    log.info("[DESKTOP SERVER V43] Iniciando em %s:%d", args.host, args.port)
    uvicorn.run(app, host=args.host, port=args.port, log_level="warning")


if __name__ == "__main__":
    main()
