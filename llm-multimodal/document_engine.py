"""
VEREDA / SYNTEXA — Document Engine
===================================
Engine de processamento de documentos com:
- PDF parsing
- DOCX parsing
- XLSX parsing
- HTML parsing
- Markdown parsing
- Layout preservation
- Metadata extraction
"""

import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass

log = logging.getLogger(__name__)


@dataclass
class DocumentResult:
    text: str
    pages: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    structure: Dict[str, Any]


class DocumentEngine:
    """
    Engine de processamento de documentos.
    """

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        deps = {}
        try:
            import fitz
            deps["PyMuPDF"] = True
        except ImportError:
            deps["PyMuPDF"] = False

        try:
            from docx import Document
            deps["python-docx"] = True
        except ImportError:
            deps["python-docx"] = False

        try:
            import openpyxl
            deps["openpyxl"] = True
        except ImportError:
            deps["openpyxl"] = False

        self._deps = deps
        log.info("Document engine deps: %s", deps)

    # ── PDF PARSING ──────────────────────────────────────────
    def parse_pdf(self, pdf_data: bytes) -> DocumentResult:
        """Parse PDF com preservação de layout."""
        if not self._deps.get("PyMuPDF"):
            return self._fallback_result("PyMuPDF não instalado")

        try:
            import fitz
            doc = fitz.open(stream=pdf_data, filetype="pdf")

            pages = []
            full_text = []
            tables = []
            images = []

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text with layout
                text = page.get_text("text")
                blocks = page.get_text("blocks")

                pages.append({
                    "page_num": page_num + 1,
                    "text": text,
                    "blocks": len(blocks),
                    "width": page.rect.width,
                    "height": page.rect.height,
                })
                full_text.append(text)

                # Extract images
                img_list = page.get_images()
                for img_index, img in enumerate(img_list, start=1):
                    images.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "xref": img[0],
                    })

            metadata = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "pages": len(doc),
                "format": "PDF",
            }

            return DocumentResult(
                text="\n\n".join(full_text),
                pages=pages,
                metadata=metadata,
                tables=tables,
                images=images,
                structure={"headings": self._extract_headings("\n".join(full_text))},
            )

        except Exception as e:
            log.error("PDF parsing failed: %s", e)
            return self._fallback_result(str(e))

    # ── DOCX PARSING ─────────────────────────────────────────
    def parse_docx(self, docx_data: bytes) -> DocumentResult:
        """Parse DOCX."""
        if not self._deps.get("python-docx"):
            return self._fallback_result("python-docx não instalado")

        try:
            from docx import Document
            from io import BytesIO

            doc = Document(BytesIO(docx_data))

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                    })

            # Extract tables
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                tables.append({"rows": rows})

            full_text = "\n".join(p["text"] for p in paragraphs)

            return DocumentResult(
                text=full_text,
                pages=[{"text": full_text, "paragraphs": len(paragraphs)}],
                metadata={"format": "DOCX", "paragraphs": len(paragraphs)},
                tables=tables,
                images=[],
                structure={"headings": [p["text"] for p in paragraphs if p["style"].startswith("Heading")]},
            )

        except Exception as e:
            log.error("DOCX parsing failed: %s", e)
            return self._fallback_result(str(e))

    # ── XLSX PARSING ─────────────────────────────────────────
    def parse_xlsx(self, xlsx_data: bytes) -> DocumentResult:
        """Parse XLSX."""
        if not self._deps.get("openpyxl"):
            return self._fallback_result("openpyxl não instalado")

        try:
            from openpyxl import load_workbook
            from io import BytesIO

            wb = load_workbook(BytesIO(xlsx_data))
            sheets = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    rows.append(list(row))
                sheets.append({"name": sheet_name, "rows": rows})

            # Convert to text
            text_parts = []
            for sheet in sheets:
                text_parts.append(f"--- {sheet['name']} ---")
                for row in sheet["rows"][:20]:  # Limit rows
                    text_parts.append(" | ".join(str(c) for c in row if c is not None))

            return DocumentResult(
                text="\n".join(text_parts),
                pages=[{"sheets": len(sheets)}],
                metadata={"format": "XLSX", "sheets": len(sheets)},
                tables=sheets,
                images=[],
                structure={"sheet_names": wb.sheetnames},
            )

        except Exception as e:
            log.error("XLSX parsing failed: %s", e)
            return self._fallback_result(str(e))

    # ── MARKDOWN / HTML ──────────────────────────────────────
    def parse_markdown(self, md_text: str) -> DocumentResult:
        """Parse Markdown."""
        import re

        headings = re.findall(r'^(#{1,6})\s+(.+)$', md_text, re.MULTILINE)
        code_blocks = re.findall(r'```(\w+)?\n(.*?)```', md_text, re.DOTALL)
        links = re.findall(r'\[(.+?)\]\((.+?)\)', md_text)

        return DocumentResult(
            text=md_text,
            pages=[{"text": md_text}],
            metadata={"format": "Markdown", "headings": len(headings)},
            tables=[],
            images=[],
            structure={
                "headings": [h[1] for h in headings],
                "code_blocks": len(code_blocks),
                "links": len(links),
            },
        )

    def parse_html(self, html_text: str) -> DocumentResult:
        """Parse HTML simplificado."""
        try:
            from html.parser import HTMLParser

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                    self.skip = False

                def handle_starttag(self, tag, attrs):
                    self.skip = tag in ("script", "style")

                def handle_endtag(self, tag):
                    self.skip = False

                def handle_data(self, data):
                    if not self.skip:
                        self.text.append(data.strip())

            extractor = TextExtractor()
            extractor.feed(html_text)

            return DocumentResult(
                text=" ".join(extractor.text),
                pages=[{"text": " ".join(extractor.text)}],
                metadata={"format": "HTML"},
                tables=[],
                images=[],
                structure={},
            )

        except Exception as e:
            return self._fallback_result(str(e))

    # ── AUTO DETECT ──────────────────────────────────────────
    def parse(self, data: bytes, content_type: Optional[str] = None) -> DocumentResult:
        """Auto-detecta formato e faz parse."""
        if content_type:
            if "pdf" in content_type:
                return self.parse_pdf(data)
            elif "wordprocessingml" in content_type or "officedocument.word" in content_type:
                return self.parse_docx(data)
            elif "spreadsheet" in content_type or "officedocument.sheet" in content_type:
                return self.parse_xlsx(data)
            elif "html" in content_type:
                return self.parse_html(data.decode('utf-8', errors='ignore'))

        # Fallback: detect by magic bytes
        if data[:4] == b'%PDF':
            return self.parse_pdf(data)
        elif data[:4] == b'PK\x03\x04':
            # Could be DOCX or XLSX
            return self.parse_docx(data)

        # Try as text
        try:
            text = data.decode('utf-8', errors='ignore')
            return self.parse_markdown(text)
        except Exception:
            return self._fallback_result("Formato não reconhecido")

    def _extract_headings(self, text: str) -> List[str]:
        import re
        lines = text.split('\n')
        headings = []
        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and not line.startswith('  ') and line[0].isupper():
                if line.endswith(':') or len(line.split()) < 8:
                    headings.append(line)
        return headings[:20]

    def _fallback_result(self, error: str) -> DocumentResult:
        return DocumentResult(
            text=f"Error: {error}",
            pages=[],
            metadata={"error": error},
            tables=[],
            images=[],
            structure={},
        )
