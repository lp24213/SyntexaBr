"""Documentos Word (python-docx) — texto + tabelas reais quando `table_rows` vier no JSON."""
from __future__ import annotations

import io
from typing import Any, Dict, List, Sequence

from docx import Document
from docx.shared import Pt


def build_docx_bytes(title: str, sections: Sequence[Dict[str, Any]]) -> bytes:
    doc = Document()
    h = doc.add_heading(title or "Documento", level=0)
    h.runs[0].font.size = Pt(18)

    for sec in sections:
        head = str(sec.get("heading") or "")
        body = str(sec.get("body") or "")
        if head:
            doc.add_heading(head, level=1)
        for para in body.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

        rows_tab = sec.get("table_rows")
        if isinstance(rows_tab, list) and rows_tab and isinstance(rows_tab[0], (list, tuple)):
            nrows = len(rows_tab)
            ncols = max(len(list(r)) for r in rows_tab) if rows_tab else 0
            if ncols > 0 and nrows > 0:
                t = doc.add_table(rows=nrows, cols=ncols)
                try:
                    t.style = "Table Grid"
                except Exception:
                    pass
                for ri, row in enumerate(rows_tab):
                    cells = list(row)
                    while len(cells) < ncols:
                        cells.append("")
                    for ci in range(ncols):
                        t.rows[ri].cells[ci].text = str(cells[ci] if ci < len(cells) else "")
                # Cópia em texto (pesquisável; leitores/Mammoth por vezes não expõem w:tbl como fluxo contínuo).
                try:
                    lines: List[str] = []
                    for row in rows_tab:
                        cells = list(row)
                        while len(cells) < ncols:
                            cells.append("")
                        lines.append("\t".join(str(cells[i]) for i in range(ncols)))
                    doc.add_paragraph("\n".join(lines)[:50000])
                except Exception:
                    pass

    try:
        footer = doc.sections[-1].footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = "© Syntexa — Todos os direitos reservados"
        if fp.runs:
            fp.runs[0].font.size = Pt(9)
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
